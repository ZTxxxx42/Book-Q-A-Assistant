"""书籍加载与分块。

支持 PDF / TXT / Markdown / EPUB / DOCX。统一输出纯文本，再交给 LightRAG 的 chunking。
"""
from __future__ import annotations

from pathlib import Path

from config import DATA_DIR, settings


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)
    return "\n\n".join(parts)


def _read_epub(path: Path) -> str:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(str(path), options={"ignore_ncx": True})
    parts: list[str] = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        parts.append(soup.get_text(separator="\n"))
    return "\n\n".join(parts)


def _read_docx(path: Path) -> str:
    # python-docx 抽段落文本。表格按行序列化为制表符分隔句行，保留结构信息。
    import docx

    doc = docx.Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n\n".join(parts)


def _read_plain(path: Path) -> str:
    # 自动尝试常见编码
    for enc in ("utf-8", "gbk", "gb18030", "utf-16"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"无法解码文件：{path}")


_READERS = {
    ".pdf": _read_pdf,
    ".epub": _read_epub,
    ".docx": _read_docx,
    ".txt": _read_plain,
    ".md": _read_plain,
    ".markdown": _read_plain,
}


def load_book(path: str | Path) -> str:
    """加载一本书，返回纯文本。"""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"书籍不存在：{path}")

    reader = _READERS.get(path.suffix.lower())
    if reader is None:
        raise ValueError(f"不支持的格式：{path.suffix}（支持 pdf/txt/md/epub/docx）")

    text = reader(path)
    text = _normalize(text)
    if not text.strip():
        raise RuntimeError(f"书籍内容为空：{path}")
    return text


def _normalize(text: str) -> str:
    """去除多余空白行，保留段落结构。"""
    import re

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_into_chunks(
    text: str,
    chunk_size: int | None = None,
    overlap: int | None = None,
) -> list[str]:
    """按字符长度切块（中文友好，不依赖分词器）。"""
    chunk_size = chunk_size or settings.chunk_size
    overlap = overlap or settings.chunk_overlap
    if chunk_size <= 0:
        return [text]

    chunks: list[str] = []
    start = 0
    step = chunk_size - overlap
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start += step
    return chunks


def resolve_book_path(name: str) -> Path:
    """data/books 下的文件名 → 完整路径。"""
    p = Path(name)
    if not p.is_absolute():
        p = DATA_DIR / name
    return p
